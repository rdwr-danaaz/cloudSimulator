#!/usr/bin/env python3
"""Generate the Word installation guide (INSTALL_GUIDE.docx) for the SOC-X simulator.

Run:
    pip install python-docx
    python deploy/generate_install_guide.py
The document is written to the repository root as INSTALL_GUIDE.docx.
"""
from __future__ import annotations

import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT = os.path.join(ROOT, "INSTALL_GUIDE.docx")

RADWARE_BLUE = RGBColor(0x1F, 0x4E, 0x79)
CODE_BG_HINT = "  "  # code paragraphs use monospace style


def add_code(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)
    p.paragraph_format.left_indent = Pt(12)
    p.paragraph_format.space_after = Pt(6)


def add_bullet(doc: Document, text: str, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)


def h(doc: Document, text: str, level: int = 1) -> None:
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RADWARE_BLUE


def build() -> None:
    doc = Document()

    # Base style
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    # ---- Title page ----
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("SOC-X Recommendation Simulator")
    tr.bold = True
    tr.font.size = Pt(26)
    tr.font.color.rgb = RADWARE_BLUE

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("Installation & Deployment Guide")
    sr.font.size = Pt(15)
    sr.italic = True

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Automated, plug-and-play installer for any Docker + ADE host").font.size = Pt(11)

    doc.add_paragraph()

    # ---- Overview ----
    h(doc, "1. Overview", 1)
    doc.add_paragraph(
        "The SOC-X Recommendation Simulator is a lightweight FastAPI service, packaged "
        "as a Docker container, that emulates the Radware SOC-X cloud recommendation API "
        "(the /_getRecommendation endpoint). It lets the Anomaly Detection Engine (ADE) "
        "receive deterministic firewall-rule recommendations in lab and staging "
        "environments without reaching the real cloud."
    )
    doc.add_paragraph(
        "The installer is fully automated. From a single command it deploys the "
        "simulator and wires it into ADE, including TLS trust and an ADE restart, so the "
        "integration works immediately."
    )
    h(doc, "What the installer does automatically", 2)
    for txt in [
        "Connects to the target machine over SSH.",
        "Auto-detects the ADE container, its Docker network, its Java truststore and its ade.config.properties file.",
        "Uploads the app, generates a stable self-signed TLS certificate, builds the image and starts the container on ADE's Docker network.",
        "Sets socx.positive.cloud.hostname and socx.remediation.cloud.hostname in ade.config.properties (a timestamped backup is taken first).",
        "Imports the simulator certificate into ADE's Java truststore so HTTPS is trusted.",
        "Restarts the ADE container so the new configuration takes effect.",
        "Verifies health and a sample recommendation from inside the ADE container.",
    ]:
        add_bullet(doc, txt)

    # ---- Architecture ----
    h(doc, "2. How it works", 1)
    doc.add_paragraph(
        "ADE always calls the SOC-X cloud over HTTPS and validates the server "
        "certificate. It also runs on its own Docker network and cannot reach arbitrary "
        "host IPs. The installer therefore:"
    )
    add_bullet(doc, "Runs the simulator on the SAME Docker network as ADE, so ADE reaches it by container name.", "Networking: ")
    add_bullet(doc, "Serves HTTPS using a self-signed certificate whose SAN matches the container name.", "TLS: ")
    add_bullet(doc, "Imports that certificate into ADE's Java truststore so validation succeeds.", "Trust: ")
    doc.add_paragraph(
        "As a result ADE is configured to call, for example, "
        "https://socx-sim:8080/api/sdcc/genai/core/analysis/peacetime/_getRecommendation."
    )

    # ---- Prerequisites ----
    h(doc, "3. Prerequisites", 1)
    h(doc, "On the machine you run the installer FROM", 2)
    add_bullet(doc, "Python 3.9 or newer.")
    add_bullet(doc, "The paramiko package:  pip install paramiko")
    add_bullet(doc, "Network/SSH access to the target machine.")
    h(doc, "On the TARGET machine", 2)
    add_bullet(doc, "Docker installed and running.")
    add_bullet(doc, "A running ADE (anomaly-detection-engine) container.")
    add_bullet(doc, "openssl available (used once to generate the TLS certificate).")

    # ---- Configuration ----
    h(doc, "4. Configuration", 1)
    doc.add_paragraph(
        "All machine-specific settings live in deploy/install_config.json. Edit this "
        "file (or pass overrides on the command line). Leaving a value blank triggers "
        "auto-detection where supported."
    )
    add_code(doc,
        '{\n'
        '  "ssh_host": "10.205.50.10",\n'
        '  "ssh_port": 22,\n'
        '  "ssh_user": "root",\n'
        '  "ssh_password": "radware",\n'
        '  "ssh_key_file": null,\n'
        '  "remote_dir": "/root/socx-sim",\n'
        '  "container": "socx-sim",\n'
        '  "host_port": 8088,\n'
        '  "internal_port": 8080,\n'
        '  "ade_container": "",            // blank = auto-detect\n'
        '  "ade_container_match": "anomaly-detection-engine",\n'
        '  "ade_network": "",             // blank = auto-detect\n'
        '  "ade_properties_path": "",     // blank = auto-detect\n'
        '  "truststore_password": "changeit",\n'
        '  "socx_hostname": "",           // blank = <container>:<internal_port>\n'
        '  "restart_ade": true,\n'
        '  "regenerate_cert": false\n'
        '}'
    )
    h(doc, "Key settings", 2)
    add_bullet(doc, "SSH connection details for the target machine.", "ssh_host / ssh_user / ssh_password / ssh_key_file: ")
    add_bullet(doc, "Name and published host port of the simulator container.", "container / host_port: ")
    add_bullet(doc, "Leave blank to auto-detect; set explicitly to override.", "ade_container / ade_network / ade_properties_path: ")
    add_bullet(doc, "The value written to ADE. Defaults to <container>:<internal_port> (e.g. socx-sim:8080).", "socx_hostname: ")
    add_bullet(doc, "Set false if you prefer to restart ADE manually.", "restart_ade: ")

    # ---- Install ----
    h(doc, "5. Installation", 1)
    doc.add_paragraph("From the repository root, run:")
    add_code(doc, "pip install paramiko\npython deploy/install.py")
    doc.add_paragraph("Common variations:")
    add_code(doc,
        "# Use a custom config file\n"
        "python deploy/install.py --config my_site.json\n\n"
        "# Override individual settings on the command line\n"
        "python deploy/install.py --ssh-host 10.0.0.5 --ssh-password secret\n\n"
        "# Deploy without restarting ADE\n"
        "python deploy/install.py --no-restart-ade\n\n"
        "# Force a fresh TLS certificate\n"
        "python deploy/install.py --regenerate-cert"
    )
    doc.add_paragraph(
        "The installer prints each step and finishes with an INSTALLATION COMPLETE "
        "summary showing the simulator URL and the address ADE uses."
    )

    # ---- Verify ----
    h(doc, "6. Verification", 1)
    doc.add_paragraph("Re-run the built-in checks at any time (no redeploy):")
    add_code(doc, "python deploy/install.py --verify-only")
    doc.add_paragraph("A healthy result looks like:")
    add_code(doc,
        "== Verifying from inside ADE ==\n"
        "  https://socx-sim:8080/health -> HTTP 200\n"
        "  recommendation sample: {\"account_id\":\"...\",\"rules\":[{...\n"
        "  RESULT: OK"
    )
    doc.add_paragraph(
        "You can also confirm from the ADE logs that recommendation calls succeed "
        "(no 'connection timed out', 'closed prematurely' or 'PKIX path building failed')."
    )

    # ---- Adding networks ----
    h(doc, "7. Adding / editing recommendation responses", 1)
    doc.add_paragraph(
        "Fixed responses are keyed by network CIDR in permanent_responses.py. To add a "
        "network, add an entry to PERMANENT_NETWORK_RULES with the CIDR as the key and "
        "the list of rule dictionaries as the value. The response timestamp and interval "
        "are computed automatically by the server, so you never edit those by hand."
    )
    doc.add_paragraph("After editing, redeploy with the same installer command:")
    add_code(doc, "python deploy/install.py")

    # ---- Uninstall ----
    h(doc, "8. Uninstall", 1)
    doc.add_paragraph("To remove the simulator and revert the ADE configuration:")
    add_code(doc, "python deploy/install.py --uninstall")
    doc.add_paragraph(
        "This stops and removes the container and image, comments out the two "
        "socx.*.cloud.hostname lines in ade.config.properties, removes the certificate "
        "from ADE's truststore and restarts ADE."
    )

    # ---- Troubleshooting ----
    h(doc, "9. Troubleshooting", 1)
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Light Grid Accent 1"
    hdr = tbl.rows[0].cells
    hdr[0].paragraphs[0].add_run("Symptom").bold = True
    hdr[1].paragraphs[0].add_run("Cause / Fix").bold = True
    rows = [
        ("ADE log: connection timed out",
         "The simulator is not on ADE's Docker network. Re-run the installer; it places the container on the detected ADE network."),
        ("ADE log: closed prematurely",
         "The simulator is serving HTTP instead of HTTPS. Re-run the installer (the image serves TLS)."),
        ("ADE log: PKIX path building failed",
         "ADE does not trust the certificate. Re-run the installer to re-import it. This is required after the ADE container is recreated."),
        ("could not find the ADE container",
         "ADE is not running, or its name does not match. Start ADE, or set ade_container in the config."),
        ("openssl: not found (target)",
         "Install openssl on the target, or place a certs/server.crt and certs/server.key in the repo before running."),
    ]
    for a, b in rows:
        cells = tbl.add_row().cells
        cells[0].paragraphs[0].add_run(a)
        cells[1].paragraphs[0].add_run(b)

    doc.add_paragraph()
    note = doc.add_paragraph()
    nr = note.add_run(
        "Note: The certificate trust and Docker-network attachment persist across "
        "'docker restart'. If the ADE container is fully recreated (e.g. docker compose "
        "down/up), simply re-run 'python deploy/install.py' to re-establish trust."
    )
    nr.italic = True

    # ---- File reference ----
    h(doc, "10. File reference", 1)
    files = [
        ("deploy/install.py", "The plug-and-play installer / uninstaller / verifier."),
        ("deploy/install_config.json", "Machine-specific settings."),
        ("cloud_mock_server.py", "The FastAPI simulator application."),
        ("permanent_responses.py", "Fixed rule responses keyed by network CIDR."),
        ("main.py", "ASGI entrypoint (serves / and /socx_sim)."),
        ("Dockerfile", "Container image; serves HTTPS using certs/server.*."),
        ("certs/", "TLS certificate and key (generated per target)."),
        ("INSTALL_GUIDE.docx", "This document."),
    ]
    for name, desc in files:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(name + " - ").bold = True
        p.add_run(desc)

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build()

